"""Turn an uploaded resume file into plain text, and text back into downloads.

Everything here is per-request and in-memory: the web app never writes a user's
resume to disk, so multiple people can use the same deployment safely.
"""

from __future__ import annotations

import io
import re
from typing import IO

SUPPORTED = ("pdf", "docx", "txt", "md")


def _clean(text: str) -> str:
    """Collapse the ragged whitespace PDF extraction leaves behind."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _from_pdf(stream: IO[bytes]) -> str:
    from pypdf import PdfReader

    reader = PdfReader(stream)
    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def _from_docx(stream: IO[bytes]) -> str:
    import docx

    doc = docx.Document(stream)
    parts = [p.text for p in doc.paragraphs]
    # Tables are common in resume templates and are invisible to `paragraphs`.
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def extract_text(uploaded_file) -> str:
    """Extract text from a Streamlit UploadedFile (or any name+bytes object).

    Raises ValueError for unsupported extensions or unreadable files.
    """
    name = getattr(uploaded_file, "name", "") or ""
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""

    # Streamlit hands back the same file object across reruns, and a stream that
    # has already been read returns b"". Rewind so a second click re-reads it
    # instead of reporting an empty file.
    if hasattr(uploaded_file, "seek"):
        try:
            uploaded_file.seek(0)
        except (OSError, ValueError):
            pass
    data = uploaded_file.read() if hasattr(uploaded_file, "read") else bytes(uploaded_file)

    if ext not in SUPPORTED:
        raise ValueError(f"Unsupported file type '.{ext}'. Upload one of: {', '.join(SUPPORTED)}.")

    try:
        if ext == "pdf":
            text = _from_pdf(io.BytesIO(data))
        elif ext == "docx":
            text = _from_docx(io.BytesIO(data))
        else:
            text = data.decode("utf-8", errors="replace")
    except Exception as exc:  # noqa: BLE001 — surfaced to the user as a friendly message
        raise ValueError(f"Could not read {name}: {exc}") from exc

    text = _clean(text)
    if len(text) < 50:
        raise ValueError(
            f"Only extracted {len(text)} characters from {name}. "
            "If it's a scanned/image PDF, paste the text manually instead."
        )
    return text


def to_docx_bytes(text: str, title: str = "Resume") -> bytes:
    """Render plain/markdown-ish text into a simple .docx for download."""
    import docx

    doc = docx.Document()
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
